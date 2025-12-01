"""
Word document generator for strategy report and statement of work.

Creates professional Word documents from synthesized deliverables.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..config import (
    OUTPUT_DIR,
    DELIVERABLES,
    SOW_BASE_PRICING,
    SOW_PRICING_MULTIPLIERS,
    CompanySize,
)
from ..models import (
    SynthesisOutput,
    ResearchOutput,
    CompanyInput,
)


class DocxGenerator:
    """
    Generates Word documents from deliverables.

    Responsibilities:
    - Create Final AI Strategy Report
    - Create Statement of Work / Engagement Letter
    - Apply consistent styling and formatting
    - Include relevant content from all deliverables
    """

    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize the Word document generator.

        Args:
            output_dir: Base output directory.
        """
        self.output_dir = output_dir or OUTPUT_DIR

    def generate_strategy_report(
        self,
        company_slug: str,
        company_input: CompanyInput,
        research: ResearchOutput,
        synthesis: SynthesisOutput,
    ) -> str:
        """
        Generate final AI strategy report.

        Comprehensive document combining all deliverables into
        a professional consulting report.

        Args:
            company_slug: URL-safe company name.
            company_input: Original company input.
            research: Research output.
            synthesis: Synthesis output.

        Returns:
            Path to generated DOCX file.
        """
        doc = Document()

        # Set up custom styles
        self._setup_styles(doc)

        # Title page
        self._add_title_page(
            doc,
            company_input.name,
            "AI Strategy & Implementation Report",
            datetime.now().strftime("%B %Y"),
        )

        # Table of contents placeholder
        self._add_table_of_contents(doc)

        # Executive summary
        self._add_section(doc, "Executive Summary", 1)
        self._add_executive_summary_content(doc, company_input, research, synthesis)

        # Company overview
        self._add_section(doc, "Company Overview", 1)
        self._add_company_overview_content(doc, research)

        # Current state assessment
        self._add_section(doc, "Current State Assessment", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "01_tech_inventory", "Technology Inventory", 2)
        self._add_subsection_from_deliverable(doc, synthesis, "02_pain_points", "Pain Point Analysis", 2)

        # AI maturity assessment
        self._add_section(doc, "AI Maturity Assessment", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "04_maturity_assessment", "Readiness Evaluation", 2)

        # Use case library
        self._add_section(doc, "AI Use Cases & Opportunities", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "14_use_case_library", "Department-Specific Use Cases", 2)

        # Strategic recommendations
        self._add_section(doc, "Strategic Recommendations", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "06_quick_wins", "Quick Wins", 2)
        self._add_subsection_from_deliverable(doc, synthesis, "05_roadmap", "Implementation Roadmap", 2)
        self._add_subsection_from_deliverable(doc, synthesis, "07_vendor_comparison", "Vendor Analysis", 2)

        # Financial analysis
        self._add_section(doc, "Financial Analysis", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "09_roi_calculator", "ROI Analysis", 2)
        self._add_subsection_from_deliverable(doc, synthesis, "08_license_consolidation", "License Consolidation", 2)

        # Governance & operations
        self._add_section(doc, "Governance & Operations", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "10_ai_policy", "AI Policy Framework", 2)
        self._add_subsection_from_deliverable(doc, synthesis, "11_data_governance", "Data Governance", 2)

        # Change management
        self._add_section(doc, "Change Management", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "15_change_management", "Training & Adoption Plan", 2)

        # Appendices
        self._add_section(doc, "Appendices", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "12_prompt_library", "Appendix A: Prompt Library", 2)
        self._add_subsection_from_deliverable(doc, synthesis, "13_glossary", "Appendix B: Glossary", 2)

        # Save document
        output_path = self._get_output_path(company_slug, "final_strategy_report.docx")
        doc.save(output_path)

        return str(output_path)

    def generate_statement_of_work(
        self,
        company_slug: str,
        company_input: CompanyInput,
        research: ResearchOutput,
        synthesis: SynthesisOutput,
    ) -> str:
        """
        Generate statement of work / engagement letter.

        Professional SOW with pricing based on company size.

        Args:
            company_slug: URL-safe company name.
            company_input: Original company input.
            research: Research output.
            synthesis: Synthesis output.

        Returns:
            Path to generated DOCX file.
        """
        doc = Document()
        self._setup_styles(doc)

        # Determine company size for pricing
        company_size = self._determine_company_size(company_input, research)

        # Title page
        self._add_title_page(
            doc,
            company_input.name,
            "Statement of Work",
            f"AI Strategy Engagement | {datetime.now().strftime('%B %Y')}",
        )

        # Introduction
        self._add_section(doc, "1. Introduction", 1)
        self._add_sow_introduction(doc, company_input)

        # Scope of work
        self._add_section(doc, "2. Scope of Work", 1)
        self._add_sow_scope(doc, synthesis)

        # Deliverables
        self._add_section(doc, "3. Deliverables", 1)
        self._add_sow_deliverables(doc)

        # Timeline
        self._add_section(doc, "4. Project Timeline", 1)
        self._add_sow_timeline(doc, synthesis)

        # Pricing
        self._add_section(doc, "5. Investment", 1)
        self._add_sow_pricing(doc, company_size)

        # Terms & conditions
        self._add_section(doc, "6. Terms & Conditions", 1)
        self._add_sow_terms(doc)

        # Signatures
        self._add_section(doc, "7. Agreement", 1)
        self._add_sow_signatures(doc, company_input)

        # Save document
        output_path = self._get_output_path(company_slug, "statement_of_work.docx")
        doc.save(output_path)

        return str(output_path)

    def _get_output_path(self, company_slug: str, filename: str) -> Path:
        """Get output path for a document."""
        output_dir = self.output_dir / company_slug / "documents"
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir / filename

    def _setup_styles(self, doc: Document) -> None:
        """Set up custom styles for the document."""
        # Get existing styles
        styles = doc.styles

        # Customize Heading 1
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = 'Arial'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(31, 78, 121)

        # Customize Heading 2
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = 'Arial'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(68, 114, 196)

        # Customize Normal
        if 'Normal' in styles:
            normal = styles['Normal']
            normal.font.name = 'Arial'
            normal.font.size = Pt(11)
            normal.font.color.rgb = RGBColor(64, 64, 64)

    def _add_title_page(
        self,
        doc: Document,
        company_name: str,
        title: str,
        subtitle: str,
    ) -> None:
        """Add title page to document."""
        # Add spacing at top
        for _ in range(8):
            doc.add_paragraph()

        # Company name
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company_para.add_run(company_name)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)

        doc.add_paragraph()

        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(64, 64, 64)

        doc.add_paragraph()

        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Add page break
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Add table of contents placeholder."""
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph(
            "[Table of contents - Update field to generate in Word]"
        )
        doc.add_page_break()

    def _add_section(self, doc: Document, title: str, level: int) -> None:
        """Add a section heading."""
        doc.add_heading(title, level=level)

    def _add_paragraph(
        self,
        doc: Document,
        text: str,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """Add a paragraph with optional formatting."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic

    def _add_bullet_list(self, doc: Document, items: List[str]) -> None:
        """Add a bullet list."""
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    def _add_table(
        self,
        doc: Document,
        headers: List[str],
        rows: List[List[str]],
    ) -> None:
        """Add a table to the document."""
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Style header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = str(value)

        doc.add_paragraph()  # Add spacing after table

    def _add_subsection_from_deliverable(
        self,
        doc: Document,
        synthesis: SynthesisOutput,
        deliverable_id: str,
        title: str,
        level: int,
    ) -> None:
        """Add subsection from a deliverable's content."""
        doc.add_heading(title, level=level)

        if deliverable_id not in synthesis.deliverables:
            doc.add_paragraph("[Content not available]")
            return

        content = synthesis.deliverables[deliverable_id].content
        if not content:
            doc.add_paragraph("[Content not available]")
            return

        # Parse markdown content and convert to Word
        self._convert_markdown_to_docx(doc, content)

    def _convert_markdown_to_docx(self, doc: Document, markdown_content: str) -> None:
        """Convert markdown content to Word format."""
        lines = markdown_content.split('\n')
        in_code_block = False
        in_table = False
        in_yaml_front_matter = False
        table_lines = []

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Handle YAML front matter (at the very start)
            if i == 0 and stripped == '---':
                in_yaml_front_matter = True
                i += 1
                continue

            if in_yaml_front_matter:
                if stripped == '---':
                    in_yaml_front_matter = False
                i += 1
                continue

            # Skip horizontal rules (---, ***, ___) - these cause the dashed lines issue
            if re.match(r'^[-*_]{3,}\s*$', stripped):
                i += 1
                continue

            # Handle code blocks
            if stripped.startswith('```'):
                # If we're starting a code block, check if it's mermaid (skip entirely)
                if not in_code_block and 'mermaid' in stripped.lower():
                    # Skip the entire mermaid block
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        i += 1
                    i += 1  # Skip closing ```
                    continue
                in_code_block = not in_code_block
                i += 1
                continue

            if in_code_block:
                # Add as monospace (but limit line length)
                para = doc.add_paragraph()
                code_line = line[:200] if len(line) > 200 else line  # Truncate very long lines
                run = para.add_run(code_line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                i += 1
                continue

            # Handle tables
            if stripped.startswith('|') and stripped.endswith('|'):
                # Skip lines that are too long (corrupted tables)
                if len(stripped) > 1000:
                    # If we were in a table, end it and skip this corrupted line
                    if in_table and table_lines:
                        self._process_markdown_table(doc, table_lines)
                        table_lines = []
                    in_table = False
                    i += 1
                    continue
                in_table = True
                table_lines.append(stripped)
                i += 1
                continue
            elif in_table:
                # End of table, process it
                if table_lines:
                    self._process_markdown_table(doc, table_lines)
                    table_lines = []
                in_table = False
                # Don't increment i - process this line normally

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Handle headings (but skip top-level as we add our own sections)
            heading_match = re.match(r'^(#{2,6})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                # Remove any remaining markdown formatting from heading
                text = self._clean_markdown_text(text)
                # Map markdown levels to Word levels (offset by 1)
                word_level = min(level, 4)
                doc.add_heading(text, level=word_level)
                i += 1
                continue

            # Handle bullet points
            bullet_match = re.match(r'^[-*+]\s+(.+)$', stripped)
            if bullet_match:
                text = self._clean_markdown_text(bullet_match.group(1))
                self._add_formatted_paragraph(doc, text, style='List Bullet')
                i += 1
                continue

            # Handle numbered lists
            numbered_match = re.match(r'^\d+\.\s+(.+)$', stripped)
            if numbered_match:
                text = self._clean_markdown_text(numbered_match.group(1))
                self._add_formatted_paragraph(doc, text, style='List Number')
                i += 1
                continue

            # Regular paragraph - add with inline formatting
            if stripped:
                # Skip excessively long lines (corrupted data)
                if len(stripped) > 2000:
                    i += 1
                    continue
                # Skip lines that are mostly dashes (table separators that weren't caught)
                if stripped.count('-') > 50 and stripped.count('-') / len(stripped) > 0.5:
                    i += 1
                    continue
                self._add_formatted_paragraph(doc, stripped)

            i += 1

        # Process any remaining table
        if table_lines:
            self._process_markdown_table(doc, table_lines)

    def _clean_markdown_text(self, text: str) -> str:
        """Remove markdown formatting markers from text."""
        # Remove markdown links but keep the text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Don't remove bold/italic markers here - we'll handle them in _add_formatted_paragraph
        return text.strip()

    def _add_formatted_paragraph(self, doc: Document, text: str, style: Optional[str] = None) -> None:
        """Add a paragraph with inline formatting (bold, italic, code)."""
        if style:
            para = doc.add_paragraph(style=style)
        else:
            para = doc.add_paragraph()

        # Parse inline formatting
        # Pattern to match **bold**, *italic*, and `code`
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Regular text
                para.add_run(part)

    def _process_markdown_table(self, doc: Document, table_lines: List[str]) -> None:
        """Process markdown table lines into a Word table."""
        if len(table_lines) < 3:
            return

        # Check if table is corrupted (any line is excessively long - indicates garbage data)
        for line in table_lines:
            if len(line) > 1000:  # Normal tables shouldn't have lines this long
                # Skip this entire corrupted table
                return

        # Parse header
        raw_headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
        # Clean headers and truncate if too long
        headers = [self._clean_table_cell(h, max_length=100) for h in raw_headers]

        # Skip if headers look corrupted (all dashes or very long)
        if all(h.startswith('-') or len(h) > 200 for h in raw_headers):
            return

        # Check if any header cell is excessively long (indicates corruption)
        if any(len(h) > 500 for h in raw_headers):
            return

        # Parse rows (skip separator line at index 1)
        rows = []
        for line in table_lines[2:]:
            # Skip separator lines (contain only dashes and colons)
            if re.match(r'^[\|\s:\-]+$', line):
                continue

            # Skip lines that are mostly dashes (corrupted data)
            dash_count = line.count('-')
            if dash_count > 100:  # Too many dashes indicates garbage
                continue

            cells = [c.strip() for c in line.strip('|').split('|')]
            if len(cells) == len(headers):
                # Check if any cell is excessively long (indicates corruption)
                if any(len(c) > 2000 for c in cells):
                    continue

                # Clean and truncate each cell
                clean_cells = [self._clean_table_cell(c, max_length=500) for c in cells]
                # Skip rows that look corrupted (mostly dashes)
                if not all(c.startswith('-') for c in cells if c):
                    # Also skip if cleaned cells are mostly empty or dashes
                    if any(len(c) > 3 and not c.replace('-', '').replace('.', '').strip() == '' for c in clean_cells):
                        rows.append(clean_cells)

        if headers and rows:
            self._add_table(doc, headers, rows)

    def _clean_table_cell(self, text: str, max_length: int = 500) -> str:
        """Clean and truncate table cell content."""
        # Remove markdown formatting
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **bold**
        text = re.sub(r'\*([^*]+)\*', r'\1', text)  # *italic*
        text = re.sub(r'`([^`]+)`', r'\1', text)  # `code`
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # [link](url)

        # Truncate if too long
        if len(text) > max_length:
            text = text[:max_length-3] + '...'

        return text.strip()

    def _add_executive_summary_content(
        self,
        doc: Document,
        company_input: CompanyInput,
        research: ResearchOutput,
        synthesis: SynthesisOutput,
    ) -> None:
        """Add executive summary content."""
        doc.add_paragraph(
            f"This report presents a comprehensive AI strategy assessment for {company_input.name}, "
            f"including current state analysis, maturity evaluation, use case identification, "
            f"and implementation recommendations."
        )

        doc.add_paragraph()

        key_findings = [
            f"Industry Position: {research.industry.primary_industry or 'Technology-focused organization'}",
            f"AI Maturity Level: {research.tech_landscape.ai_maturity_estimate or 'Early exploration phase'}",
            "Multiple high-impact use cases identified across departments",
            "Quick wins available for immediate value realization",
            "Comprehensive 360-day roadmap for transformation",
        ]

        doc.add_heading("Key Findings", level=2)
        self._add_bullet_list(doc, key_findings)

    def _add_company_overview_content(
        self,
        doc: Document,
        research: ResearchOutput,
    ) -> None:
        """Add company overview content."""
        profile = research.profile

        if profile.description:
            doc.add_paragraph(profile.description)

        doc.add_heading("Company Profile", level=2)

        profile_items = []
        if profile.business_model:
            profile_items.append(f"Business Model: {profile.business_model}")
        if profile.company_size:
            profile_items.append(f"Company Size: {profile.company_size.value.title()}")
        if profile.headquarters:
            profile_items.append(f"Headquarters: {profile.headquarters}")
        if profile.products_services:
            profile_items.append(f"Key Products/Services: {', '.join(profile.products_services[:5])}")

        if profile_items:
            self._add_bullet_list(doc, profile_items)

        # Industry context
        industry = research.industry
        if industry.primary_industry:
            doc.add_heading("Industry Context", level=2)
            doc.add_paragraph(f"Primary Industry: {industry.primary_industry}")

            if industry.key_trends:
                doc.add_paragraph("Key Industry Trends:", style='Normal')
                self._add_bullet_list(doc, industry.key_trends[:5])

    # =========================================================================
    # SOW-specific methods
    # =========================================================================

    def _determine_company_size(
        self,
        company_input: CompanyInput,
        research: ResearchOutput,
    ) -> CompanySize:
        """Determine company size from input and research."""
        # Check explicit input first
        if company_input.employee_count:
            count = company_input.employee_count
            if count < 100:
                return CompanySize.SMALL
            elif count < 500:
                return CompanySize.MEDIUM
            elif count < 2000:
                return CompanySize.LARGE
            else:
                return CompanySize.ENTERPRISE

        # Check research data
        if research.profile.company_size:
            return research.profile.company_size

        if research.profile.employee_estimate:
            count = research.profile.employee_estimate
            if count < 100:
                return CompanySize.SMALL
            elif count < 500:
                return CompanySize.MEDIUM
            elif count < 2000:
                return CompanySize.LARGE
            else:
                return CompanySize.ENTERPRISE

        # Default to medium
        return CompanySize.MEDIUM

    def _add_sow_introduction(
        self,
        doc: Document,
        company_input: CompanyInput,
    ) -> None:
        """Add SOW introduction section."""
        doc.add_paragraph(
            f"This Statement of Work ('SOW') defines the scope, deliverables, timeline, "
            f"and investment for AI strategy consulting services to be provided to "
            f"{company_input.name} ('Client')."
        )

        doc.add_paragraph()

        doc.add_heading("Objectives", level=2)
        objectives = [
            "Assess current AI readiness and maturity",
            "Identify high-impact AI use cases aligned with business goals",
            "Develop comprehensive implementation roadmap",
            "Provide governance and change management frameworks",
            "Enable sustainable AI adoption and transformation",
        ]
        self._add_bullet_list(doc, objectives)

    def _add_sow_scope(self, doc: Document, synthesis: SynthesisOutput) -> None:
        """Add SOW scope section."""
        phases = [
            ("Discovery & Assessment", [
                "Technology inventory and infrastructure assessment",
                "Pain point analysis by department",
                "AI maturity evaluation",
                "Stakeholder interviews and requirements gathering",
            ]),
            ("Strategy Development", [
                "Use case identification and prioritization",
                "Quick wins analysis",
                "Vendor evaluation and build vs buy assessment",
                "ROI modeling and financial analysis",
            ]),
            ("Roadmap & Planning", [
                "30/60/90/180/360-day implementation roadmap",
                "Resource planning and capability requirements",
                "Risk assessment and mitigation strategies",
                "Success metrics and KPI definition",
            ]),
            ("Governance & Enablement", [
                "AI policy and acceptable use framework",
                "Data governance guidelines",
                "Change management and training plan",
                "Prompt library and best practices",
            ]),
        ]

        for phase_name, activities in phases:
            doc.add_heading(phase_name, level=2)
            self._add_bullet_list(doc, activities)

    def _add_sow_deliverables(self, doc: Document) -> None:
        """Add SOW deliverables section."""
        doc.add_paragraph(
            "The following deliverables will be provided upon completion of the engagement:"
        )

        deliverables = [
            ("Assessment Reports", [
                "Technology Inventory & Data Infrastructure Assessment",
                "Pain Point Matrix by Department",
                "AI Maturity Model & Readiness Assessment",
            ]),
            ("Strategy Documents", [
                "Department-Specific Use Case Library",
                "Quick Wins List with Implementation Guides",
                "30/60/90/180/360 Implementation Roadmap",
                "Vendor Comparison & Build vs Buy Analysis",
            ]),
            ("Financial Analysis", [
                "ROI Calculator & Cost-Benefit Analysis",
                "License Consolidation Recommendations",
            ]),
            ("Governance & Operations", [
                "AI Acceptable Use Policy Template",
                "Data Governance Framework",
                "Change Management & Training Playbook",
            ]),
            ("Supporting Materials", [
                "Prompt Library Starter Kit",
                "Glossary of AI Terms",
                "Current & Future State Architecture Diagrams",
            ]),
            ("Presentations", [
                "Executive Summary Deck",
                "Full Findings & Recommendations Presentation",
            ]),
        ]

        for category, items in deliverables:
            doc.add_heading(category, level=2)
            self._add_bullet_list(doc, items)

    def _add_sow_timeline(self, doc: Document, synthesis: SynthesisOutput) -> None:
        """Add SOW timeline section."""
        headers = ["Phase", "Duration", "Key Activities"]
        rows = [
            ["Discovery", "Week 1-2", "Assessments, interviews, data collection"],
            ["Analysis", "Week 3-4", "Maturity assessment, use case identification"],
            ["Strategy", "Week 5-6", "Roadmap development, vendor evaluation"],
            ["Documentation", "Week 7-8", "Report writing, presentation preparation"],
            ["Review & Delivery", "Week 9-10", "Final reviews, client presentation, handoff"],
        ]

        doc.add_paragraph("Estimated project duration: 8-10 weeks")
        doc.add_paragraph()
        self._add_table(doc, headers, rows)

    def _add_sow_pricing(self, doc: Document, company_size: CompanySize) -> None:
        """Add SOW pricing section."""
        multiplier = SOW_PRICING_MULTIPLIERS.get(company_size, 1.0)

        if multiplier is None:
            # Enterprise - custom pricing
            doc.add_paragraph(
                "For enterprise-scale engagements, pricing is customized based on "
                "scope, complexity, and specific requirements. Please contact us "
                "to discuss your needs."
            )
            return

        # Calculate prices
        discovery = int(SOW_BASE_PRICING["discovery"] * multiplier)
        strategy = int(SOW_BASE_PRICING["strategy"] * multiplier)
        implementation = int(SOW_BASE_PRICING["implementation_support"] * multiplier)
        training = int(SOW_BASE_PRICING["training"] * multiplier)
        total = discovery + strategy + implementation + training

        headers = ["Component", "Description", "Investment"]
        rows = [
            ["Discovery & Assessment", "Technology, maturity, and pain point analysis", f"${discovery:,}"],
            ["Strategy Development", "Use cases, roadmap, and vendor analysis", f"${strategy:,}"],
            ["Implementation Support", "Governance frameworks and documentation", f"${implementation:,}"],
            ["Training & Enablement", "Change management and knowledge transfer", f"${training:,}"],
            ["Total Investment", "", f"${total:,}"],
        ]

        self._add_table(doc, headers, rows)

        doc.add_paragraph()
        doc.add_heading("Payment Terms", level=2)
        payment_terms = [
            "50% due upon SOW signature",
            "25% due at mid-project milestone (Week 5)",
            "25% due upon final delivery and acceptance",
        ]
        self._add_bullet_list(doc, payment_terms)

    def _add_sow_terms(self, doc: Document) -> None:
        """Add SOW terms and conditions."""
        terms = [
            "Confidentiality: All information shared during the engagement will be treated as confidential.",
            "Intellectual Property: Client retains ownership of all deliverables upon final payment.",
            "Changes: Any changes to scope will be documented in a change order with adjusted timeline and pricing.",
            "Cancellation: Either party may terminate with 14 days written notice. Client pays for work completed.",
            "Liability: Consultant liability limited to fees paid under this agreement.",
            "Warranty: Deliverables warranted to meet specifications for 30 days after delivery.",
        ]

        for term in terms:
            doc.add_paragraph(term)
            doc.add_paragraph()

    def _add_sow_signatures(
        self,
        doc: Document,
        company_input: CompanyInput,
    ) -> None:
        """Add SOW signature section."""
        doc.add_paragraph(
            "By signing below, both parties agree to the terms outlined in this Statement of Work."
        )

        doc.add_paragraph()

        # Client signature block
        doc.add_heading("Client", level=2)
        doc.add_paragraph(f"Company: {company_input.name}")
        doc.add_paragraph("Name: _______________________________")
        doc.add_paragraph("Title: _______________________________")
        doc.add_paragraph("Signature: _______________________________")
        doc.add_paragraph("Date: _______________________________")

        doc.add_paragraph()

        # Consultant signature block
        doc.add_heading("Consultant", level=2)
        doc.add_paragraph("Company: AI Strategy Factory")
        doc.add_paragraph("Name: _______________________________")
        doc.add_paragraph("Title: _______________________________")
        doc.add_paragraph("Signature: _______________________________")
        doc.add_paragraph("Date: _______________________________")


# Convenience functions

def generate_strategy_report(
    company_slug: str,
    company_input: CompanyInput,
    research: ResearchOutput,
    synthesis: SynthesisOutput,
    output_dir: Optional[Path] = None,
) -> str:
    """Generate final AI strategy report."""
    generator = DocxGenerator(output_dir=output_dir)
    return generator.generate_strategy_report(
        company_slug, company_input, research, synthesis
    )


def generate_statement_of_work(
    company_slug: str,
    company_input: CompanyInput,
    research: ResearchOutput,
    synthesis: SynthesisOutput,
    output_dir: Optional[Path] = None,
) -> str:
    """Generate statement of work."""
    generator = DocxGenerator(output_dir=output_dir)
    return generator.generate_statement_of_work(
        company_slug, company_input, research, synthesis
    )
